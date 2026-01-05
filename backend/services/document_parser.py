"""
Document parser service for extracting text from PDF and DOCX files
"""
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """Custom exception for document parsing errors"""
    pass


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text content from a PDF file
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        DocumentParseError: If file cannot be parsed or is corrupted
        FileNotFoundError: If file does not exist
        ValueError: If file is empty or has no readable content
    """
    # Check if file exists
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    # Check if file is empty
    if file_path.stat().st_size == 0:
        raise ValueError(f"PDF file is empty: {file_path}")
    
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            # Check if PDF has pages
            if len(pdf.pages) == 0:
                raise ValueError(f"PDF file has no pages: {file_path}")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
            
            # Check if any text was extracted
            if not text.strip():
                raise ValueError(f"PDF file contains no extractable text: {file_path}")
                
    except FileNotFoundError:
        raise
    except ValueError:
        raise
    except pdfplumber.PDFSyntaxError as e:
        raise DocumentParseError(f"PDF file appears to be corrupted or invalid: {str(e)}")
    except Exception as e:
        raise DocumentParseError(f"Error extracting text from PDF: {str(e)}")
    
    return text.strip()


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text content from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        DocumentParseError: If file cannot be parsed or is corrupted
        FileNotFoundError: If file does not exist
        ValueError: If file is empty or has no readable content
    """
    # Check if file exists
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    # Check if file is empty
    if file_path.stat().st_size == 0:
        raise ValueError(f"DOCX file is empty: {file_path}")
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Also try to extract text from tables if present
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs_text + tables_text
        
        # Check if any text was extracted
        if not all_text:
            raise ValueError(f"DOCX file contains no extractable text: {file_path}")
        
        text = "\n".join(all_text)
        
    except FileNotFoundError:
        raise
    except ValueError:
        raise
    except Exception as e:
        # Check for common DOCX errors
        error_msg = str(e).lower()
        if "not a zip file" in error_msg or "bad zipfile" in error_msg:
            raise DocumentParseError(f"DOCX file appears to be corrupted or invalid (not a valid ZIP archive): {file_path}")
        elif "cannot open" in error_msg or "invalid" in error_msg:
            raise DocumentParseError(f"DOCX file cannot be opened or is invalid: {str(e)}")
        else:
            raise DocumentParseError(f"Error extracting text from DOCX: {str(e)}")
    
    return text.strip()


def parse_document(file_path: Path) -> str:
    """
    Parse a document and extract text based on file extension
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If file extension is not supported
        FileNotFoundError: If file does not exist
        DocumentParseError: If file cannot be parsed
    """
    # Validate file path
    if not isinstance(file_path, Path):
        file_path = Path(file_path)
    
    file_extension = file_path.suffix.lower()
    
    # Check if extension is supported
    if file_extension not in [".pdf", ".docx"]:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx")
    
    # Route to appropriate parser
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
